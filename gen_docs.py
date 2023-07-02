import docx
import pkg_resources
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor


class generate_dependency_docx():

    def __init__(self) -> None:
        self.document_title = 'Project Dependencies'
        self.header_font_size = 13
        self.table_header_font_family = 'Calibri'
        self.column_headers = [
            "Package",
            "Version",
            "License",
            "Author",
            "Description",
            "Homepage URL"
        ]
        self.table_font_family = 'Calibri'
        self.alt_table_row_color = ['EEEEEE', 'FFFFFF']
        self.font_normal_color = RGBColor(90, 90, 90)
        self.font_highlight_color = RGBColor(0x0E, 0x8D, 0xC6)
        self.font_not_found_color = RGBColor(255, 0, 0)
        self.font_header_color = RGBColor(0, 120, 156)
        
    def get_package_license(self, package_name) -> tuple:
        """
        Retrieves the license information for a given package.

        Args:
            package_name (str): The name of the package.

        Returns:
            tuple: The license information for the package.
        """
        try:
            package_dist = pkg_resources.get_distribution(package_name)
            metadata = package_dist.get_metadata('METADATA')  # Adjust metadata file based on the package metadata format
            for line in metadata.split('\n'):
                if line.startswith('License:'):
                    return (True, line.split(':', 1)[1].strip() )
        except pkg_resources.DistributionNotFound:
            pass

        return (False, "Unknown")  # Default value if license information is not found

    def get_package_author(self, package_name) -> tuple:
        """
        Retrieves the author information for a given package.

        Args:
            package_name (str): The name of the package.

        Returns:
            tuple: The author information for the package.
        """
        try:
            package_dist = pkg_resources.get_distribution(package_name)
            metadata = package_dist.get_metadata('METADATA')  # Adjust metadata file based on the package metadata format
            for line in metadata.split('\n'):
                if line.startswith('Author:'):
                    return (True, line.split(':', 1)[1].strip() )
        except pkg_resources.DistributionNotFound:
            pass

        return (False, 'Unknown')  # Default value if author information is not found

    def get_package_description(self, package_name) -> tuple:
        """
        Retrieves the description for a given package.

        Args:
            package_name (str): The name of the package.

        Returns:
            tuple: The description of the package.
        """
        try:
            package_dist = pkg_resources.get_distribution(package_name)
            metadata = package_dist.get_metadata('METADATA')  # Adjust metadata file based on the package metadata format
            for line in metadata.split('\n'):
                if line.startswith('Summary:'):
                    return (True, line.split(':', 1)[1].strip() )
        except pkg_resources.DistributionNotFound:
            pass

        return (False, 'No description available')  # Default value if description is not found

    def get_package_homepage(self, package_name) -> tuple:
        """
        Retrieves the homepage URL for a given package.

        Args:
            package_name (str): The name of the package.

        Returns:
            tuple: The homepage URL of the package.
        """
        try:
            package_dist = pkg_resources.get_distribution(package_name)
            metadata = package_dist.get_metadata('METADATA')  # Adjust metadata file based on the package metadata format
            for line in metadata.split('\n'):
                if line.startswith('Home-page:'):
                    return (True, line.split(':', 1)[1].strip() )
        except pkg_resources.DistributionNotFound:
            pass

        return (False, 'No homepage URL available')  # Default value if homepage URL is not found

    def get_package_info(self) -> list:
        """
        Retrieves package information for all installed packages.

        Returns:
            list: A list of dictionaries representing package information.
        """
        installed_packages = pkg_resources.working_set
        package_info_list = []

        for package in installed_packages:
            package_info = {
                'name': package.key,
                'version': package.version,
                'license': self.get_package_license(package.key),
                'author': self.get_package_author(package.key),
                'description': self.get_package_description(package.key),
                'homepage_url': self.get_package_homepage(package.key),
            }
            package_info_list.append(package_info)

        return package_info_list

    def generate_report(self, dependencies) -> None:
        """
        Generates a report in DOCX format listing the project dependencies.

        Args:
            dependencies (list): A list of tuples containing the package names and their versions.

        Returns:
            None
        """
        doc = Document()
        doc.add_heading(self.document_title, level=0 )

        # Create a table for the dependencies
        num_columns = len(self.column_headers)
        num_rows = 1
        table = doc.add_table(rows=num_rows, cols=num_columns)
        table.autofit = False
        table.columns[0].width = Cm(1)
        table.columns[1].width = Cm(1)
        table.columns[2].width = Cm(1)
        table.columns[3].width = Cm(1)
        table.columns[4].width = Inches(1.5)
        table.columns[5].width = Inches(1.3)
        
        # Set table headers
        headers = table.rows[0].cells
        for index, column_name in enumerate(self.column_headers):
            headers[index].text = column_name
            headers[index].paragraphs[0].runs[0].bold = True  # Make header text bold
            headers[index].paragraphs[0].runs[0].font.size = Pt(self.header_font_size) 
            headers[index].paragraphs[0].runs[0].font.name = self.table_header_font_family 
            headers[index].paragraphs[0].runs[0].font.color.rgb = self.font_header_color

        # Populate the table with the dependencies
        for index, dependency in enumerate(dependencies):
            row_cells = table.add_row().cells
            row_cells[0].text = dependency['name']
            row_cells[0].paragraphs[0].runs[0].font.name = self.table_font_family
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = self.font_highlight_color
            
            row_cells[1].text = dependency['version']
            row_cells[1].paragraphs[0].runs[0].font.name = self.table_font_family 
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = self.font_normal_color
            
            row_cells[2].text = dependency['license'][1]
            row_cells[2].paragraphs[0].runs[0].font.name = self.table_font_family 
            row_cells[2].paragraphs[0].runs[0].font.color.rgb = self.font_normal_color if dependency['license'][0] else self.font_not_found_color
            
            row_cells[3].text = dependency['author'][1]
            row_cells[3].paragraphs[0].runs[0].font.name = self.table_font_family 
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = self.font_normal_color if dependency['author'][0] else self.font_not_found_color
            
            row_cells[4].text = dependency['description'][1]
            row_cells[4].paragraphs[0].runs[0].font.name = self.table_font_family 
            row_cells[4].paragraphs[0].runs[0].font.color.rgb = self.font_normal_color if dependency['description'][0] else self.font_not_found_color
            
            row_cells[5].text = dependency['homepage_url'][1]
            row_cells[5].paragraphs[0].runs[0].font.name = self.table_font_family 
            row_cells[5].paragraphs[0].runs[0].font.color.rgb = self.font_normal_color if dependency['homepage_url'][0] else self.font_not_found_color
            
            # Define alternating row colors
            colors = self.alt_table_row_color
            
            # Apply alternating row colors
            color_index = index % len(colors)
            for cell in row_cells:
                cell_paragraph = cell.paragraphs[0]
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="{colors[color_index]}" />'
                shading_element = parse_xml(shading_xml)
                cell._element.tcPr.append(shading_element)

        # Save the report as a DOCX file
        doc.save('dependencies_report.docx')

        print('Report generated successfully.')


if __name__ == '__main__':
    # dependencies = get_dependencies()
    document = generate_dependency_docx()
    dependencies = document.get_package_info()
    document.generate_report(dependencies)
